from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware

from parser.logic import analyze_scene
from parser.splitter import split_scenes

from docx import Document
from io import BytesIO
import pdfplumber
from striprtf.striprtf import rtf_to_text
import re

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================
# 📌 DOCX ПОЛНОСТЬЮ СЛИВАЕТ ВСЕ RUN'ы В ЕДИНУЮ ЛИНИЮ
# ============================================================

def extract_docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))

    lines = []

    for p in doc.paragraphs:
        full = "".join(run.text for run in p.runs)
        full = full.replace("\u00A0", " ").strip()

        # Word может делать сцена № как буквы по одной → склеиваем вручную
        full = full.replace("С Ц Е Н А", "СЦЕНА")
        full = full.replace("С Ц Е Н А ", "СЦЕНА ")
        full = full.replace("С  Ц  Е  Н  А", "СЦЕНА")

        if full:
            lines.append(full)

    # Таблицы Word
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = "".join(run.text for run in p.runs)
                    full = full.replace("\u00A0", " ").strip()

                    full = full.replace("С Ц Е Н А", "СЦЕНА")
                    full = full.replace("С Ц Е Н А ", "СЦЕНА ")

                    if full:
                        lines.append(full)

    text = "\n".join(lines)

    # Нормализация
    text = text.replace("\u00A0", " ")
    text = text.replace("\u2028", "\n")
    text = re.sub(r"\n+", "\n", text)
    text = re.sub(r"[ ]+", " ", text)

    return text.strip()


# ============================================================
# 📌 PDF
# ============================================================

def extract_pdf_text(content: bytes) -> str:
    result = ""

    with pdfplumber.open(BytesIO(content)) as pdf:
        for page in pdf.pages:
            tx = page.extract_text()
            if not tx:
                continue

            tx = tx.replace("-\n", "")
            tx = tx.replace("\n", " ")

            result += tx + " "

    result = result.replace("\u00A0", " ")
    return " ".join(result.split())


# ============================================================
# 📌 Общий экстрактор
# ============================================================

def extract_text_from_file(filename: str, content: bytes) -> str:
    ext = filename.lower().split(".")[-1]

    if ext == "txt":
        try:
            return content.decode("utf-8")
        except:
            return content.decode("cp1251", errors="ignore")

    if ext == "docx":
        return extract_docx_text(content)

    if ext == "pdf":
        return extract_pdf_text(content)

    if ext == "rtf":
        try:
            return rtf_to_text(content.decode("utf-8", errors="ignore"))
        except:
            return rtf_to_text(content.decode("cp1251", errors="ignore"))

    return ""


# ============================================================
# 📌 API
# ============================================================

@app.post("/parse_file")
async def parse_file(file: UploadFile = File(...)):
    content = await file.read()
    text = extract_text_from_file(file.filename, content)



    scenes = split_scenes(text)
    result = []

    for header, body in scenes:
        result.append({
            "scene_header": header,
            "analysis": analyze_scene(body)
        })

    return {"scenes": result}


@app.post("/parse_text")
async def parse_text(data: dict):
    text = data.get("text", "")
    scenes = split_scenes(text)

    result = []
    for header, body in scenes:
        result.append({
            "scene_header": header,
            "analysis": analyze_scene(body)
        })

    return {"scenes": result}


if __name__ == "__main__":
    import uvicorn
    print(">>> http://127.0.0.1:8000")
    uvicorn.run(app, host="127.0.0.1", port=8000)
